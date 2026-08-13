import json
import os
import tempfile
import time
import urllib.error
import urllib.request
import wave
from difflib import SequenceMatcher
from io import BytesIO
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from threading import Event, Lock, Thread
from urllib.parse import urlparse

from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt
import numpy as np

from audio_capture import capture_backend, capture_capabilities, get_desktop_capture_microphone
from batch_capture import BatchManager


PROJECT_DIR = Path(__file__).resolve().parent
PORT = int(os.getenv("ASR_PORT", "3211"))
MODEL_NAME = os.getenv("LOCAL_ASR_MODEL", "medium")
DEVICE = os.getenv("LOCAL_ASR_DEVICE", "cpu")
COMPUTE_TYPE = os.getenv("LOCAL_ASR_COMPUTE_TYPE", "int8")
MAX_AUDIO_BYTES = int(os.getenv("MAX_AUDIO_BYTES", str(25 * 1024 * 1024)))
MAX_JSON_BYTES = 2 * 1024 * 1024


def load_env():
    env_path = PROJECT_DIR / ".env"
    if not env_path.exists():
        return
    for raw_line in env_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("#") and "OPENAI_API_KEY=" in line:
            line = line[line.index("OPENAI_API_KEY="):]
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


load_env()
AI_API_KEY = os.getenv("OPENAI_API_KEY", "")
AI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1").rstrip("/")
AI_MODEL = os.getenv("AI_MODEL", "gpt-4.1-mini")

_model = None
_model_lock = Lock()


class DesktopRecorder:
    def __init__(self):
        self.lock = Lock()
        self.stop_event = Event()
        self.thread = None
        self.status = "idle"
        self.error = None
        self.started_at = None
        self.frames = 0
        self.sample_rate = 48000
        self.channels = 2
        self.file_path = None
        self.output_name = None

    def snapshot(self):
        with self.lock:
            duration = self.frames / self.sample_rate if self.sample_rate else 0
            return {
                "status": self.status,
                "error": self.error,
                "startedAt": self.started_at,
                "duration": round(duration, 1),
                "outputName": self.output_name,
                "canRetranscribe": bool(self.file_path and self.file_path.exists() and self.status in {"complete", "error"}),
            }

    def start(self, title="desktop-audio"):
        with self.lock:
            if self.status in {"starting", "recording", "stopping", "processing"}:
                raise ValueError("已有桌面录音任务正在进行。")
            safe_title = "".join("-" if char in '\\/:*?\"<>|' else char for char in str(title or "desktop-audio"))
            safe_title = " ".join(safe_title.split())[:80] or "desktop-audio"
            timestamp = time.strftime("%Y%m%d-%H%M%S")
            self.output_name = f"{safe_title}-{timestamp}.wav"
            recordings_dir = PROJECT_DIR / "recordings"
            recordings_dir.mkdir(exist_ok=True)
            previous_path = self.file_path
            if previous_path and previous_path.exists():
                previous_path.unlink(missing_ok=True)
            self.file_path = recordings_dir / self.output_name
            self.stop_event.clear()
            self.status = "starting"
            self.error = None
            self.started_at = int(time.time() * 1000)
            self.frames = 0
            self.thread = Thread(target=self._record_loop, daemon=True)
            self.thread.start()
        return self.snapshot()

    def _record_loop(self):
        try:
            microphone = get_desktop_capture_microphone()
            with wave.open(str(self.file_path), "wb") as wav_file:
                wav_file.setnchannels(self.channels)
                wav_file.setsampwidth(2)
                wav_file.setframerate(self.sample_rate)
                with microphone.recorder(samplerate=self.sample_rate, channels=self.channels) as recorder:
                    with self.lock:
                        self.status = "recording"
                    while not self.stop_event.is_set():
                        data = recorder.record(numframes=4800)
                        if data is None or not len(data):
                            continue
                        pcm = (np.clip(data, -1.0, 1.0) * 32767).astype(np.int16)
                        wav_file.writeframes(pcm.tobytes())
                        with self.lock:
                            self.frames += len(pcm)
        except Exception as error:
            with self.lock:
                self.status = "error"
                self.error = str(error) or "桌面音频录制失败。"

    def stop_and_transcribe(self):
        with self.lock:
            if self.status not in {"starting", "recording"}:
                raise ValueError(self.error or "当前没有正在进行的桌面录音。")
            self.status = "stopping"
            thread = self.thread
            file_path = self.file_path
            output_name = self.output_name
        self.stop_event.set()
        if thread:
            thread.join(timeout=10)
        if thread and thread.is_alive():
            raise RuntimeError("停止桌面录音超时，请重启本地服务。")
        with self.lock:
            if self.error:
                raise RuntimeError(self.error)
            if self.frames < self.sample_rate:
                raise ValueError("录音不足1秒，请播放内容后再停止。")
            self.status = "processing"
        try:
            result = transcribe(file_path, vad_filter=False)
            if not result["segments"]:
                raise ValueError("没有识别到有效语音，请确认红果 App 正在播放且系统音量正常。")
            result["filename"] = output_name
            result["source"] = capture_backend()
            with self.lock:
                self.status = "complete"
            return result
        except Exception as error:
            with self.lock:
                self.status = "error"
                self.error = str(error) or "桌面音频识别失败。"
            raise

    def retranscribe(self):
        with self.lock:
            if self.status in {"starting", "recording", "stopping", "processing"}:
                raise ValueError("当前任务尚未结束，请稍后再重新识别。")
            file_path = self.file_path
            output_name = self.output_name
            if not file_path or not file_path.exists():
                raise ValueError("最近一次桌面录音未保留，请重新录制一遍。")
            self.status = "processing"
            self.error = None
        try:
            result = transcribe(file_path, vad_filter=False)
            if not result["segments"]:
                raise ValueError("完整模式仍未识别到有效语音，请检查系统音量。")
            result["filename"] = output_name
            result["source"] = "desktop-loopback-full"
            with self.lock:
                self.status = "complete"
            return result
        except Exception as error:
            with self.lock:
                self.status = "error"
                self.error = str(error) or "完整模式重新识别失败。"
            raise

    def cancel(self):
        with self.lock:
            if self.status not in {"starting", "recording", "stopping"}:
                raise ValueError("当前没有可以取消的桌面录音。")
            thread = self.thread
            file_path = self.file_path
        self.stop_event.set()
        if thread:
            thread.join(timeout=10)
        if thread and thread.is_alive():
            raise RuntimeError("取消桌面录音超时，请重启本地服务。")
        if file_path:
            file_path.unlink(missing_ok=True)
        with self.lock:
            self.status = "idle"
            self.error = None
            self.started_at = None
            self.frames = 0
            self.thread = None
            self.file_path = None
            self.output_name = None
        return self.snapshot()


_desktop_recorder = DesktopRecorder()


def get_model():
    global _model
    if _model is None:
        with _model_lock:
            if _model is None:
                print(
                    f"Loading faster-whisper model: {MODEL_NAME} "
                    f"({DEVICE}/{COMPUTE_TYPE})",
                    flush=True,
                )
                _model = WhisperModel(
                    MODEL_NAME,
                    device=DEVICE,
                    compute_type=COMPUTE_TYPE,
                    download_root=str(PROJECT_DIR / "models"),
                )
                print("Model loaded.", flush=True)
    return _model


def transcribe(audio_path, vad_filter=False):
    model = get_model()
    segment_stream, info = model.transcribe(
        str(audio_path),
        language="zh",
        beam_size=5,
        vad_filter=vad_filter,
        vad_parameters={"min_silence_duration_ms": 250} if vad_filter else None,
        condition_on_previous_text=True,
    )

    segments = []
    for index, segment in enumerate(segment_stream):
        text = segment.text.strip()
        if not text:
            continue
        segments.append(
            {
                "id": index,
                "start": round(float(segment.start), 3),
                "end": round(float(segment.end), 3),
                "text": text,
            }
        )

    return {
        "model": f"faster-whisper/{MODEL_NAME}",
        "language": info.language or "zh",
        "duration": round(float(info.duration or (segments[-1]["end"] if segments else 0)), 3),
        "text": "".join(item["text"] for item in segments),
        "segments": segments,
    }


def timeline_text(segments):
    def seconds_value(value):
        if isinstance(value, str) and ":" in value:
            parts = value.strip().split(":")
            try:
                total = 0.0
                for part in parts:
                    total = total * 60 + float(part)
                return total
            except ValueError:
                return 0.0
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    def stamp(value):
        seconds = max(0, int(seconds_value(value)))
        return f"{seconds // 60:02d}:{seconds % 60:02d}"
    return "\n".join(
        f"[{stamp(item.get('start'))}-{stamp(item.get('end'))}] "
        f"{(str(item.get('speaker', '')).strip() + '：') if str(item.get('speaker', '')).strip() else ''}"
        f"{str(item.get('text', '')).strip()}"
        for item in segments
        if str(item.get("text", "")).strip()
    )


def extract_response_text(payload):
    if isinstance(payload.get("output_text"), str):
        return payload["output_text"]
    choices = payload.get("choices") or []
    if choices:
        content = choices[0].get("message", {}).get("content", "")
        if isinstance(content, str):
            return content
    parts = []
    for output in payload.get("output") or []:
        for content in output.get("content") or []:
            if isinstance(content.get("text"), str):
                parts.append(content["text"])
    return "\n".join(parts)


def post_ai(url, payload):
    request = urllib.request.Request(
        url,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {AI_API_KEY}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=180) as response:
        return json.loads(response.read().decode("utf-8"))


def call_ai(system_prompt, user_prompt):
    if not AI_API_KEY:
        raise ValueError("未配置文本模型密钥，无法生成AI拆解。")

    attempts = [
        (
            f"{AI_BASE_URL}/responses",
            {"model": AI_MODEL, "instructions": system_prompt, "input": user_prompt, "max_output_tokens": 5000},
        ),
        (
            f"{AI_BASE_URL}/v1/responses",
            {"model": AI_MODEL, "instructions": system_prompt, "input": user_prompt, "max_output_tokens": 5000},
        ),
        (
            f"{AI_BASE_URL}/chat/completions",
            {"model": AI_MODEL, "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]},
        ),
        (
            f"{AI_BASE_URL}/v1/chat/completions",
            {"model": AI_MODEL, "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]},
        ),
    ]
    errors = []
    for url, payload in attempts:
        try:
            result = post_ai(url, payload)
            text = extract_response_text(result).strip()
            if text:
                return text
            errors.append(f"{url}: 返回空内容")
        except urllib.error.HTTPError as error:
            errors.append(f"{url}: HTTP {error.code}")
        except Exception as error:
            errors.append(f"{url}: {error}")
    raise ValueError("中转站文本接口不可用：" + "；".join(errors))


def parse_ai_json(raw):
    value = raw.strip()
    if value.startswith("```"):
        value = value.split("\n", 1)[1].rsplit("```", 1)[0]
    return json.loads(value)


def analyze_project(payload):
    segments = payload.get("segments") or []
    if not segments:
        raise ValueError("没有可分析的时间轴文本。")
    title = str(payload.get("title") or "未命名短剧").strip()
    episode = str(payload.get("episode") or "第1集").strip()
    prompt = f"""剧名：{title}\n集数：{episode}\n\n校对后的时间轴：\n{timeline_text(segments)}\n\n请严格输出JSON对象，不要Markdown代码块。"""
    system = """你是中文短剧编剧顾问。请先结合上下文修正明显的同音错字，并根据人物关系、上下句问答和对白中的称谓为每句推断说话角色。时间轴中已经人工填写且不等于“待确认”的speaker是人工确认结果，必须原样保留；只允许为空白或“待确认”的speaker补充角色。优先使用对白可支持的姓名或称谓，例如“妞妞”“掌教”“师伯”“姐姐”“众长老”；只能判断身份时可写“女主”“母亲”“长老”，完全无法判断时speaker写“待确认”，不得强行编造姓名。每个id、时间和原顺序必须保持不变，不得把相邻台词错位。再进行原创创作分析，不复刻原作表达。输出JSON，字段必须为：corrected_segments（数组，元素含id/start/end/speaker/text）、summary（字符串）、episode_function（字符串，说明本集在整剧中的叙事功能，例如引入人物、建立冲突、升级矛盾、承上启下或制造付费悬念）、characters（字符串数组）、opening_hook（字符串）、conflict_nodes（字符串数组）、emotion_curve（字符串，按照开场→发展→高潮→结尾描述本集观众情绪如何变化）、reversals（字符串数组）、ending_hook（字符串）、creative_methods（字符串数组）、original_ideas（字符串数组）。不得虚构时间轴中不存在的关键剧情。"""
    result = parse_ai_json(call_ai(system, prompt))
    corrected = result.get("corrected_segments") or []
    corrected_by_id = {str(item.get("id")): item for item in corrected if isinstance(item, dict)}
    normalized_corrections = []
    for index, original in enumerate(segments):
        correction = corrected_by_id.get(str(original.get("id")))
        if correction is None and len(corrected) == len(segments) and index < len(corrected):
            correction = corrected[index]
        original_text = str(original.get("text") or "").strip()
        proposed_text = str((correction or {}).get("text") or original_text).strip()
        similarity = SequenceMatcher(None, original_text, proposed_text).ratio() if original_text and proposed_text else 1
        safe_text = proposed_text if similarity >= 0.45 else original_text
        original_speaker = str(original.get("speaker") or "").strip()
        proposed_speaker = str((correction or {}).get("speaker") or "").strip()
        safe_speaker = original_speaker if original_speaker and original_speaker != "待确认" else (proposed_speaker or original_speaker or "待确认")
        normalized_corrections.append(
            {
                "id": original.get("id", index),
                "start": original.get("start", 0),
                "end": original.get("end", original.get("start", 0)),
                "speaker": safe_speaker,
                "text": safe_text,
            }
        )
    result["corrected_segments"] = normalized_corrections
    result["title"] = title
    result["episode"] = episode
    return result


def episode_digest(episode):
    analysis = episode.get("analysis") or {}
    segments = episode.get("segments") or []
    return {
        "episode": episode.get("episode") or episode.get("label") or "未知集数",
        "summary": analysis.get("summary") or "".join(str(item.get("text", "")) for item in segments)[:1200],
        "episode_function": analysis.get("episode_function") or "",
        "characters": analysis.get("characters") or [],
        "opening_hook": analysis.get("opening_hook") or "",
        "conflict_nodes": analysis.get("conflict_nodes") or [],
        "emotion_curve": analysis.get("emotion_curve") or analysis.get("emotion_beats") or "",
        "reversals": analysis.get("reversals") or [],
        "ending_hook": analysis.get("ending_hook") or "",
    }


def summarize_episode_group(title, episodes):
    system = """你是短剧长篇结构分析师。根据连续若干集的标准化单集资料，输出严格JSON，不要代码块。字段：range（字符串）、stage_summary（字符串）、character_changes（字符串数组）、conflict_progression（字符串数组）、hook_chain（字符串数组）、emotion_curve（字符串）、key_turning_points（字符串数组）。不得补写资料中不存在的情节。"""
    payload = {"title": title, "episodes": [episode_digest(item) for item in episodes]}
    return parse_ai_json(call_ai(system, json.dumps(payload, ensure_ascii=False)))


def analyze_series(payload):
    title = str(payload.get("title") or "未命名短剧").strip()
    scope = str(payload.get("scope") or "全部").strip()
    episodes = payload.get("episodes") or []
    if not episodes:
        raise ValueError("项目中没有可分析的集数。")

    groups = []
    for start in range(0, len(episodes), 10):
        batch = episodes[start:start + 10]
        groups.append(summarize_episode_group(title, batch))

    system = """你是资深中文短剧总编剧。根据各阶段摘要生成整剧或阶段创作分析，强调结构规律和原创迁移，不复刻原作表达。严格输出JSON，不要代码块。字段必须为：overall_summary（字符串）、structure_stages（字符串数组）、main_characters（字符串数组）、character_arcs（字符串数组）、core_conflicts（字符串数组）、hook_patterns（字符串数组）、payment_beats（字符串数组）、emotion_curve（字符串）、pacing_analysis（字符串）、hit_elements（字符串数组）、creative_formulas（字符串数组）、original_directions（字符串数组）、episode_index（字符串数组）。不得虚构资料中不存在的关键剧情。"""
    prompt = json.dumps(
        {"title": title, "scope": scope, "episode_count": len(episodes), "stage_reports": groups},
        ensure_ascii=False,
    )
    result = parse_ai_json(call_ai(system, prompt))
    result.update({"title": title, "scope": scope, "episode_count": len(episodes), "stage_reports": groups})
    return result


def build_docx(payload):
    analysis = payload.get("analysis") or {}
    segments = payload.get("segments") or []
    title = str(payload.get("title") or analysis.get("title") or "未命名短剧")
    episode = str(payload.get("episode") or analysis.get("episode") or "第1集")

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    document.add_heading("短剧创作拆解报告", 0)
    document.add_paragraph(f"剧名：{title}")
    document.add_paragraph(f"集数：{episode}")

    document.add_heading("校对后的时间轴文本", level=1)
    for line in timeline_text(segments).splitlines():
        document.add_paragraph(line)

    sections = [
        ("单集剧情梗概", analysis.get("summary", "")),
        ("本集作用", analysis.get("episode_function", "")),
        ("人物与关系", analysis.get("characters", [])),
        ("开场钩子", analysis.get("opening_hook", "")),
        ("冲突升级节点", analysis.get("conflict_nodes", [])),
        ("情绪曲线", analysis.get("emotion_curve") or analysis.get("emotion_beats", [])),
        ("反转点", analysis.get("reversals", [])),
        ("结尾悬念", analysis.get("ending_hook", "")),
        ("可借鉴的创作方法", analysis.get("creative_methods", [])),
        ("原创改编建议", analysis.get("original_ideas", [])),
    ]
    for heading, content in sections:
        document.add_heading(heading, level=1)
        if isinstance(content, list):
            for item in content:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(content or "暂无"))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_series_docx(payload):
    report = payload.get("report") or {}
    episodes = payload.get("episodes") or []
    title = str(payload.get("title") or report.get("title") or "未命名短剧")
    scope = str(payload.get("scope") or report.get("scope") or "整体分析")
    document = Document()
    document.styles["Normal"].font.name = "Microsoft YaHei"
    document.styles["Normal"].font.size = Pt(10.5)
    document.add_heading("短剧多集创作分析报告", 0)
    document.add_paragraph(f"剧名：{title}")
    document.add_paragraph(f"分析范围：{scope}")
    document.add_paragraph(f"已纳入集数：{len(episodes)}集")

    document.add_heading("分集完整资料", level=1)
    for episode in episodes:
        analysis = episode.get("analysis") or {}
        label = str(episode.get("episode") or episode.get("label") or "未知集数")
        summary = str(analysis.get("summary") or "暂无单集拆解")
        document.add_heading(label, level=2)
        document.add_heading("本集作用", level=3)
        document.add_paragraph(str(analysis.get("episode_function") or "暂无"))
        document.add_heading("单集剧情梗概", level=3)
        document.add_paragraph(summary)
        document.add_heading("校对后的角色时间轴", level=3)
        lines = timeline_text(episode.get("segments") or []).splitlines()
        if lines:
            for line in lines:
                document.add_paragraph(line)
        else:
            document.add_paragraph("暂无时间轴")
        episode_sections = [
            ("人物与关系", analysis.get("characters", [])),
            ("开场钩子", analysis.get("opening_hook", "")),
            ("冲突升级节点", analysis.get("conflict_nodes", [])),
            ("情绪曲线", analysis.get("emotion_curve") or analysis.get("emotion_beats", [])),
            ("反转点", analysis.get("reversals", [])),
            ("结尾悬念", analysis.get("ending_hook", "")),
            ("可借鉴的创作方法", analysis.get("creative_methods", [])),
            ("原创改编建议", analysis.get("original_ideas", [])),
        ]
        for subheading, content in episode_sections:
            document.add_heading(subheading, level=3)
            if isinstance(content, list):
                for item in content:
                    document.add_paragraph(str(item), style="List Bullet")
            else:
                document.add_paragraph(str(content or "暂无"))

    document.add_page_break()
    document.add_heading("多集整体创作分析", level=1)
    sections = [
        ("整体故事概述", report.get("overall_summary", "")),
        ("阶段结构", report.get("structure_stages", [])),
        ("主要人物", report.get("main_characters", [])),
        ("人物成长弧线", report.get("character_arcs", [])),
        ("核心冲突", report.get("core_conflicts", [])),
        ("钩子模式", report.get("hook_patterns", [])),
        ("付费卡点", report.get("payment_beats", [])),
        ("情绪曲线", report.get("emotion_curve", "")),
        ("节奏分析", report.get("pacing_analysis", "")),
        ("爆款元素", report.get("hit_elements", [])),
        ("可迁移创作公式", report.get("creative_formulas", [])),
        ("原创方向建议", report.get("original_directions", [])),
    ]
    for heading, content in sections:
        document.add_heading(heading, level=1)
        if isinstance(content, list):
            for item in content:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(content or "暂无"))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_batch_manager = BatchManager(
    PROJECT_DIR,
    transcribe_fn=lambda path: transcribe(path, vad_filter=False),
    analyze_fn=analyze_project,
)


class Handler(BaseHTTPRequestHandler):
    server_version = "LocalASR/1.0"

    def send_json(self, status, payload):
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, X-Audio-Filename")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.end_headers()
        self.wfile.write(body)

    def send_bytes(self, status, body, content_type, filename=None):
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        if filename:
            self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.end_headers()
        self.wfile.write(body)

    def read_json(self):
        length = int(self.headers.get("Content-Length", "0"))
        if length <= 0 or length > MAX_JSON_BYTES:
            raise ValueError("请求数据为空或过大。")
        return json.loads(self.rfile.read(length).decode("utf-8"))

    def require_local_client(self):
        origin = self.headers.get("Origin", "")
        if origin and not (
            origin.startswith("chrome-extension://")
            or origin.startswith("http://127.0.0.1")
            or origin.startswith("http://localhost")
        ):
            raise PermissionError("桌面录音接口只允许本机扩展调用。")

    def do_OPTIONS(self):
        self.send_json(204, {})

    def do_GET(self):
        route = urlparse(self.path).path
        if route == "/api/health":
            self.send_json(
                200,
                {
                    "ok": True,
                    "configured": True,
                    "engine": "faster-whisper",
                    "model": MODEL_NAME,
                    "loaded": _model is not None,
                    **capture_capabilities(),
                },
            )
            return
        if route == "/api/desktop-recording/status":
            try:
                self.require_local_client()
                self.send_json(200, _desktop_recorder.snapshot())
            except Exception as error:
                self.send_json(403, {"error": str(error)})
            return
        if route == "/api/batch/status":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.status())
            except Exception as error:
                self.send_json(403, {"error": str(error)})
            return
        if route == "/api/batch/results":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.results())
            except Exception as error:
                self.send_json(403, {"error": str(error)})
            return
        self.send_json(404, {"error": "Not found"})

    def do_POST(self):
        route = urlparse(self.path).path
        if route == "/api/batch/start":
            try:
                self.require_local_client()
                if _desktop_recorder.snapshot().get("status") in {"starting", "recording", "stopping", "processing"}:
                    raise ValueError("单集桌面录音正在进行，请先停止。")
                payload = self.read_json()
                self.send_json(200, _batch_manager.start(payload.get("title"), payload.get("startEpisode"), payload.get("endEpisode")))
            except Exception as error:
                self.send_json(500, {"error": str(error) or "连续采集启动失败。"})
            return
        if route == "/api/batch/next":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.next_episode())
            except Exception as error:
                self.send_json(500, {"error": str(error) or "切换下一集失败。"})
            return
        if route == "/api/batch/finish":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.finish())
            except Exception as error:
                self.send_json(500, {"error": str(error) or "结束连续采集失败。"})
            return
        if route == "/api/batch/cancel":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.cancel())
            except Exception as error:
                self.send_json(500, {"error": str(error) or "取消任务失败。"})
            return
        if route == "/api/batch/delete":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.delete())
            except Exception as error:
                self.send_json(500, {"error": str(error) or "删除任务失败。"})
            return
        if route == "/api/batch/retry":
            try:
                self.require_local_client()
                payload = self.read_json()
                self.send_json(200, _batch_manager.retry(payload.get("episode")))
            except Exception as error:
                self.send_json(500, {"error": str(error) or "重新识别失败。"})
            return
        if route == "/api/batch/ai/start":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.start_ai())
            except Exception as error:
                self.send_json(500, {"error": str(error) or "批量AI启动失败。"})
            return
        if route == "/api/batch/controller/open":
            try:
                self.require_local_client()
                self.send_json(200, _batch_manager.open_controller(PORT))
            except Exception as error:
                self.send_json(500, {"error": str(error) or "置顶控制器启动失败。"})
            return
        if route == "/api/desktop-recording/start":
            try:
                self.require_local_client()
                if _batch_manager.status().get("status") == "recording":
                    raise ValueError("连续多集采集正在进行，请先结束。")
                payload = self.read_json()
                self.send_json(200, _desktop_recorder.start(payload.get("title", "desktop-audio")))
            except Exception as error:
                print(f"Desktop recording start failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "桌面录音启动失败。"})
            return

        if route == "/api/desktop-recording/stop":
            try:
                self.require_local_client()
                self.send_json(200, _desktop_recorder.stop_and_transcribe())
            except Exception as error:
                print(f"Desktop recording stop failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "桌面录音停止失败。"})
            return

        if route == "/api/desktop-recording/cancel":
            try:
                self.require_local_client()
                self.send_json(200, _desktop_recorder.cancel())
            except Exception as error:
                print(f"Desktop recording cancel failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "取消桌面录音失败。"})
            return

        if route == "/api/desktop-recording/retranscribe":
            try:
                self.require_local_client()
                self.send_json(200, _desktop_recorder.retranscribe())
            except Exception as error:
                print(f"Desktop retranscription failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "完整模式重新识别失败。"})
            return

        if route == "/api/analyze":
            try:
                self.send_json(200, analyze_project(self.read_json()))
            except Exception as error:
                print(f"Analysis failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "AI拆解失败。"})
            return

        if route == "/api/analyze-series":
            try:
                self.send_json(200, analyze_series(self.read_json()))
            except Exception as error:
                print(f"Series analysis failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "多集整体分析失败。"})
            return

        if route == "/api/export-docx":
            try:
                body = build_docx(self.read_json())
                self.send_bytes(
                    200,
                    body,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "short-drama-analysis.docx",
                )
            except Exception as error:
                print(f"DOCX export failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "Word导出失败。"})
            return

        if route == "/api/export-series-docx":
            try:
                body = build_series_docx(self.read_json())
                self.send_bytes(
                    200,
                    body,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "short-drama-series-analysis.docx",
                )
            except Exception as error:
                print(f"Series DOCX export failed: {error}", flush=True)
                self.send_json(500, {"error": str(error) or "多集Word导出失败。"})
            return

        if route != "/api/transcribe":
            self.send_json(404, {"error": "Not found"})
            return

        try:
            length = int(self.headers.get("Content-Length", "0"))
            if length <= 0:
                raise ValueError("收到的音频为空。")
            if length > MAX_AUDIO_BYTES:
                raise ValueError(f"音频超过本地限制 {MAX_AUDIO_BYTES // 1024 // 1024}MB。")

            audio = self.rfile.read(length)
            suffix = Path(self.headers.get("X-Audio-Filename", "recording.webm")).suffix
            if suffix.lower() not in {".webm", ".wav", ".mp3", ".m4a", ".mp4", ".ogg"}:
                suffix = ".webm"

            temp_path = None
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                    temp_file.write(audio)
                    temp_path = Path(temp_file.name)
                result = transcribe(temp_path)
            finally:
                if temp_path:
                    temp_path.unlink(missing_ok=True)

            if not result["segments"]:
                raise ValueError("没有识别到有效语音，请确认录音音量正常。")
            self.send_json(200, result)
        except Exception as error:
            print(f"Transcription failed: {error}", flush=True)
            self.send_json(500, {"error": str(error) or "本地语音识别失败。"})

    def log_message(self, fmt, *args):
        print(f"{self.address_string()} - {fmt % args}", flush=True)


if __name__ == "__main__":
    server = ThreadingHTTPServer(("127.0.0.1", PORT), Handler)
    print(f"Local ASR service: http://127.0.0.1:{PORT}", flush=True)
    print(f"Model: faster-whisper/{MODEL_NAME}", flush=True)
    print("The model is downloaded and loaded on the first transcription.", flush=True)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()
